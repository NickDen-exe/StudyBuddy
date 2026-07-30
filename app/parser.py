import pymupdf
import docx
import os
from pathlib import Path

DOCS_PATH = "./data"


def load_docs():
    pdf_docs = []
    docx_docs = []

    if not os.access(DOCS_PATH, os.F_OK):
        print("Access to docs folder denied")
        return pdf_docs, docx_docs

    for file_path in Path(DOCS_PATH).iterdir():
        if file_path.suffix.lower() == ".pdf":
            pdf_docs.append(file_path)
        elif file_path.suffix.lower() == ".docx":
            docx_docs.append(file_path)

    if not pdf_docs and not docx_docs:
        print("No PDF or DOCX files found in the docs folder.")

    return pdf_docs, docx_docs


def read_pdf(pdf_docs):
    pdf_data = []
    for pdf_doc in pdf_docs:
        try:
            doc = pymupdf.Document(pdf_doc)
            for page in doc:
                text = page.get_text()
                pdf_data.append({"text": text, "page": page.number + 1, "source": pdf_doc.name})
            doc.close()
        except Exception as e:
            print(f"Error reading PDF file {pdf_doc.name}: {e}")
    return pdf_data


def read_docx(docx_docs):
    docx_data = []
    for docx_doc in docx_docs:
        try:
            doc = docx.Document(docx_doc)
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    docx_data.append({"text": text, "paragraph": i + 1, "source": docx_doc.name})
        except Exception as e:
            print(f"Error reading DOCX file {docx_doc.name}: {e}")
    return docx_data


def read_data():
    pdf_docs, docx_docs = load_docs()
    data = []
    data.extend(read_pdf(pdf_docs))
    data.extend(read_docx(docx_docs))
    return data